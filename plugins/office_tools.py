"""
Plugin: office_tools.py
Kemampuan membaca dan membuat dokumen produktivitas:
PDF, Word (.docx), Excel (.xlsx), dan CSV.
"""
import os

# ── PDF ───────────────────────────────────────────────────────────────────────

def read_pdf(file_path: str, max_pages: int = 10) -> str:
    """Mengekstrak teks dari file PDF."""
    try:
        import fitz  # PyMuPDF
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages_to_read = min(total_pages, max_pages)
        texts = []

        for i in range(pages_to_read):
            page = doc[i]
            text = page.get_text().strip()
            if text:
                texts.append(f"--- Halaman {i + 1} ---\n{text}")

        doc.close()

        if not texts:
            return f"PDF '{os.path.basename(file_path)}' tidak mengandung teks yang bisa dibaca (mungkin PDF berisi gambar/scan)."

        result = "\n\n".join(texts)
        note = f"\n\n[Menampilkan {pages_to_read} dari {total_pages} halaman.]" if total_pages > max_pages else ""
        return f"Isi PDF '{os.path.basename(file_path)}':\n\n{result}{note}"

    except ImportError:
        return "Library 'PyMuPDF' belum diinstal. Jalankan: pip install PyMuPDF"
    except Exception as e:
        return f"Gagal membaca PDF: {e}"

def get_pdf_info(file_path: str) -> str:
    """Mendapatkan informasi metadata PDF (judul, jumlah halaman, author, dll)."""
    try:
        import fitz
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        doc = fitz.open(file_path)
        meta = doc.metadata
        size_mb = round(os.path.getsize(file_path) / (1024 * 1024), 2)
        doc.close()

        info = [
            f"File: {os.path.basename(file_path)}",
            f"Ukuran: {size_mb} MB",
            f"Jumlah Halaman: {doc.page_count}",
            f"Judul: {meta.get('title') or '(tidak ada)'}",
            f"Author: {meta.get('author') or '(tidak ada)'}",
            f"Dibuat: {meta.get('creationDate') or '(tidak ada)'}",
            f"Dimodifikasi: {meta.get('modDate') or '(tidak ada)'}",
        ]
        return "\n".join(info)
    except Exception as e:
        return f"Gagal membaca info PDF: {e}"

# ── Word (.docx) ──────────────────────────────────────────────────────────────

def read_docx(file_path: str) -> str:
    """Mengekstrak teks dari file Word (.docx)."""
    try:
        from docx import Document
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        if not paragraphs:
            return f"Dokumen Word '{os.path.basename(file_path)}' kosong atau tidak ada teks."

        content = "\n".join(paragraphs)
        # Batasi output agar tidak terlalu panjang
        if len(content) > 6000:
            content = content[:6000] + "\n\n... (konten terpotong karena terlalu panjang)"

        return f"Isi dokumen '{os.path.basename(file_path)}':\n\n{content}"
    except ImportError:
        return "Library 'python-docx' belum diinstal. Jalankan: pip install python-docx"
    except Exception as e:
        return f"Gagal membaca file Word: {e}"

def create_docx(file_path: str, title: str, content: str) -> str:
    """Membuat file Word (.docx) baru dengan judul dan konten teks."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Tambah judul
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tambah konten (mendukung newline sebagai paragraf baru)
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()  # Baris kosong

        # Pastikan direktori ada
        os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
        doc.save(file_path)
        return f"File Word berhasil dibuat: '{os.path.abspath(file_path)}'"
    except ImportError:
        return "Library 'python-docx' belum diinstal. Jalankan: pip install python-docx"
    except Exception as e:
        return f"Gagal membuat file Word: {e}"

def get_docx_info(file_path: str) -> str:
    """Mendapatkan info singkat tentang dokumen Word (jumlah paragraf, kata)."""
    try:
        from docx import Document
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        word_count = sum(len(p.split()) for p in paragraphs)
        size_kb = round(os.path.getsize(file_path) / 1024, 1)

        return (
            f"File: {os.path.basename(file_path)}\n"
            f"Ukuran: {size_kb} KB\n"
            f"Jumlah Paragraf: {len(paragraphs)}\n"
            f"Estimasi Jumlah Kata: {word_count}"
        )
    except Exception as e:
        return f"Gagal membaca info Word: {e}"

# ── Excel (.xlsx) ─────────────────────────────────────────────────────────────

def read_excel(file_path: str, sheet_name: str = None, max_rows: int = 50) -> str:
    """Membaca data dari file Excel (.xlsx). Mengembalikan data dalam format tabel teks."""
    try:
        import openpyxl
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

        # Pilih sheet
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                return f"Sheet '{sheet_name}' tidak ditemukan. Sheet yang tersedia: {', '.join(wb.sheetnames)}"
            ws = wb[sheet_name]
        else:
            ws = wb.active
            sheet_name = ws.title

        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                rows.append(f"... (data terpotong, menampilkan {max_rows} baris pertama dari {ws.max_row} total baris)")
                break
            row_str = " | ".join([str(cell) if cell is not None else "" for cell in row])
            rows.append(row_str)

        wb.close()

        if not rows:
            return f"Sheet '{sheet_name}' kosong."

        return f"Data Excel '{os.path.basename(file_path)}' - Sheet '{sheet_name}':\n\n" + "\n".join(rows)
    except ImportError:
        return "Library 'openpyxl' belum diinstal. Jalankan: pip install openpyxl"
    except Exception as e:
        return f"Gagal membaca Excel: {e}"

def get_excel_info(file_path: str) -> str:
    """Mendapatkan informasi tentang file Excel: daftar sheet dan dimensi data."""
    try:
        import openpyxl
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        wb = openpyxl.load_workbook(file_path, read_only=True)
        info = [f"File: {os.path.basename(file_path)}"]
        info.append(f"Jumlah Sheet: {len(wb.sheetnames)}")

        for name in wb.sheetnames:
            ws = wb[name]
            info.append(f"  - Sheet '{name}': {ws.max_row} baris x {ws.max_column} kolom")

        wb.close()
        return "\n".join(info)
    except Exception as e:
        return f"Gagal membaca info Excel: {e}"

# ── CSV ───────────────────────────────────────────────────────────────────────

def read_csv(file_path: str, max_rows: int = 50) -> str:
    """Membaca dan menampilkan isi file CSV."""
    try:
        import csv
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' tidak ditemukan."

        rows = []
        with open(file_path, 'r', encoding='utf-8-sig', errors='replace') as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    rows.append(f"... (menampilkan {max_rows} baris pertama)")
                    break
                rows.append(" | ".join(row))

        if not rows:
            return f"File CSV '{os.path.basename(file_path)}' kosong."

        return f"Isi CSV '{os.path.basename(file_path)}':\n\n" + "\n".join(rows)
    except Exception as e:
        return f"Gagal membaca CSV: {e}"

# ── Tool Definitions ──────────────────────────────────────────────────────────

OFFICE_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Membaca dan mengekstrak teks dari file PDF. Mendukung multi-halaman.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path lengkap ke file PDF."},
                    "max_pages": {"type": "integer", "description": "Maksimal halaman yang dibaca (default: 10)."}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_pdf_info",
            "description": "Mendapatkan metadata PDF: jumlah halaman, judul, author, tanggal pembuatan.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path lengkap ke file PDF."}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "Membaca isi dokumen Microsoft Word (.docx) dan mengembalikan teksnya.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path lengkap ke file Word .docx."}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": "Membuat file Microsoft Word (.docx) baru dengan judul dan isi konten teks.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path output file .docx yang akan dibuat."},
                    "title": {"type": "string", "description": "Judul dokumen Word."},
                    "content": {"type": "string", "description": "Isi teks dokumen. Gunakan \\n untuk baris baru."}
                },
                "required": ["file_path", "title", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_docx_info",
            "description": "Mendapatkan info ringkas dari dokumen Word: ukuran, jumlah paragraf, estimasi kata.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path ke file .docx."}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_excel",
            "description": "Membaca data tabel dari file Microsoft Excel (.xlsx). Bisa pilih sheet tertentu.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path ke file .xlsx."},
                    "sheet_name": {"type": "string", "description": "Nama sheet yang ingin dibaca. Kosongkan untuk sheet pertama."},
                    "max_rows": {"type": "integer", "description": "Maksimal baris yang ditampilkan (default: 50)."}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_excel_info",
            "description": "Melihat daftar sheet dan dimensi (baris x kolom) dari file Excel.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path ke file .xlsx."}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_csv",
            "description": "Membaca dan menampilkan isi file CSV sebagai tabel teks.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path ke file .csv."},
                    "max_rows": {"type": "integer", "description": "Maksimal baris yang ditampilkan (default: 50)."}
                },
                "required": ["file_path"]
            }
        }
    }
]

OFFICE_TOOL_MAP = {
    "read_pdf": read_pdf,
    "get_pdf_info": get_pdf_info,
    "read_docx": read_docx,
    "create_docx": create_docx,
    "get_docx_info": get_docx_info,
    "read_excel": read_excel,
    "get_excel_info": get_excel_info,
    "read_csv": read_csv,
}
